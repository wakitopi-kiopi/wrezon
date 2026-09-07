
import schemas

from fastapi.middleware.cors import CORSMiddleware

from fastapi import FastAPI, Response,HTTPException
from fpdf import FPDF
import re
import io
import matplotlib
import matplotlib.pyplot as plt
from docx import Document
from io import BytesIO
import json
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


matplotlib.use('Agg') #to prio protect matplotlib from defaulting to server dispay engine

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["https://wrezon.netlify.app",
                "https://wrezon.onrender.com",
                "https://wrez.netlify.app",
                "https://www.wrezon.com",
                "http://localhost:8000",
                "http://127.0.0.1:5501",
                "http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"]
)


class WrezonPDF(FPDF):
  
  def __init__(self,docname):
    super().__init__()
    self.docname = docname
    self.add_font("DejaVuSans", "", "fonts/DejaVuSans.ttf")
    self.add_font("DejaVuSans", "B", "fonts/DejaVuSans-Bold.ttf")
    self.add_font("DejaVuSans", "I", "fonts/DejaVuSans-Oblique.ttf")

  def header(self):
    self.set_font("DejaVuSans", "B", 14)
    self.cell(0, 10, self.docname, border=False, new_x="LMARGIN",
        new_y="NEXT", align="C")
    self.ln(5)

  def footer(self):
    self.set_y(-15)
    self.set_font("DejaVuSans", "B", 8)
    self.cell(0, 10, f"Wrezon Export |Page {self.page_no()}/{{nb}}", align="C")

def make_math_image(formula: str) -> io.BytesIO:
    formula = formula.replace('\\\\', '\\')
    fig = plt.figure(figsize=(0.1, 0.1))
    fig.text(0, 0, f"${formula}$", fontsize=12)
    
    buf = io.BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight', pad_inches=0.02, transparent=True, dpi=300)
    plt.close(fig)
    buf.seek(0)
    return buf
  
  


def clean_markdown(text: str) -> str:
    """Remove markdown formatting characters from plain text."""
    # Remove bold **text**
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # Remove italic *text*
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    # Remove headers ## text
    text = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)
    # Remove strikethrough ~~text~~
    text = re.sub(r'~~(.*?)~~', r'\1', text)
    # Remove inline code `text`
    text = re.sub(r'`(.*?)`', r'\1', text)
    # Remove code block markers ```
    text = re.sub(r'```.*?```', '', text, flags=re.DOTALL)
    return text

@app.post("/export-pdf")
async def export_pdf(text_content: schemas.pdf_struct):
  docname = text_content.docname
  try:
    pdf = WrezonPDF(docname)
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Standard readable text formatting
    pdf.set_font("DejaVuSans", size=11)
    
    
    safe_text = text_content.query
    #.encode('latin-1', 'replace').decode('latin-1')
    items = re.split(r'(\$\$.*?\$\$|\$.*?\$)', safe_text)
    for item in items:
        if not item:
            continue
        if ((item.startswith('$$') and item.endswith('$$')) or 
            (item.startswith('$') and item.endswith('$')) or 
            (item.startswith('\\(') and item.endswith('\\)')) or 
            (item.startswith('\\[') and item.endswith('\\]'))):
            formula = item.replace('\\(', '').replace('\\)', '').replace('\\[', '').replace('\\]', '').replace('$$', '').replace('$', '')
            
            # 3. IF it's Math (starts and ends with $)
           #if item.startswith('$') and item.endswith('$'):
            formula = formula.replace('\\\\', '\\')
            formula = formula.strip('$') # Strip $ signs
            
            # Generate image bytes in RAM via Matplotlib
            math_bytes = make_math_image(formula)
            
            # Insert the image into FPDF at the current cursor position
            pdf.image(math_bytes, h=10)
            
        # 4. ELSE it's plain text
        else:
            # Print text at current cursor position
            cleaned_text = clean_markdown(item)
            pdf.write(5, cleaned_text)
            #pdf.multi_cell(0, 7,item)
    
    # multi_cell automatically handles word wrapping and margins
    

    # Stream output directly to memory (avoids writing to disk)
    #output = BytesIO()
    #pdf.output(output)
    #pdf_bytes = output.getvalue()
    pdf_bytes = bytes(pdf.output())

    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={docname}.pdf"},
    )
  
    
    
  except Exception as e:
    print("error with PDF generation",(e))
    raise HTTPException(
        status_code=500, detail=f"Failed to generate PDF: {str(e)}"
    )
    
@app.post("/export_word")
async def create_word_document(contents:schemas.wodr_struct):
    heading=contents.wdocname
    body =json.loads(contents.query)
    doctitle =body.get('title')
    sections =body.get('sections')
    
    doc = Document()
    title = doc.add_heading(doctitle, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    
    section = doc.sections[0]

    footer = section.footer
    
    paragraph = footer.paragraphs[0]
    paragraph.add_run("wrezon Export | ")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")

    paragraph._p.append(field)
    
    for entry in sections:
        section_title = entry.get('title')
        level =entry.get('level')
        content = entry.get('content')
        
        sec_head = doc.add_heading(section_title,level=level)
        sec_head.alignment =WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(content)
    
    buffer=BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{heading}.docx"'
        }
    )
 
 
 
 
 
 
    
    
@app.post("/health")
def awake():
    status = "200 OK"
    return status
  
@app.get("/")
@app.head("/")
async def root():
    return {"status": "ok"}
  
